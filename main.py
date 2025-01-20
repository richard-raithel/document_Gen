import os
import dotenv
from docx import Document
from openai import OpenAI

# Load environment variables from .env file
dotenv.load_dotenv()

# Initialize OpenAI API
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def update_resume(docx_file, job_description, employer_name, job_title, output_folder):
    """
    Updates only the text in the first table (which contains 'Experience' and 'Projects'),
    leaving other paragraphs (header) untouched.
    """
    # Load the resume
    doc = Document(docx_file)

    # 1) Identify paragraphs (header) — we'll keep them as context if you want,
    #    but won't modify them.
    header_paragraphs = [p.text for p in doc.paragraphs]

    # 2) Identify the correct table (assumes the first table is the left table
    #    with "Experience" and "Projects").
    #    If you have multiple tables, you can do logic to find the table that
    #    actually contains "Experience" or "Projects" in one of its cells.
    if len(doc.tables) == 0:
        print("No tables found in the document!")
        return

    table_to_update = doc.tables[0]  # the first (left) table

    # 3) Extract all cell text from the chosen table, with markers
    #    so we can preserve row/cell structure when re-inserting.
    table_blocks = []
    for r, row in enumerate(table_to_update.rows):
        for c, cell in enumerate(row.cells):
            marker = f"<ROW {r} CELL {c}>"
            cell_text = cell.text
            table_blocks.append(f"{marker}\n{cell_text}\n")

    # Combine table text into one chunk
    table_text_combined = "".join(table_blocks)

    # 4) Build the prompt, giving the header as read-only context if desired
    #    The LLM sees the header but is instructed NOT to modify it; it only
    #    updates the table text.
    prompt = """
        You are a resume-updating assistant.
        
        Below is a header (context only, do not modify) and then a table which should be updated.
        
        HEADER (DO NOT CHANGE):
        {}
        
        TABLE TEXT (UPDATE THIS):
        {}
        
        JOB DESCRIPTION:
        {}
        
        Please update ONLY the text in the table to better match the job description while
        retaining all rows/cells, structure, style, and paragraph breaks within each cell.
        
        Important Guidelines:
        - Do NOT remove or add rows/cells.
        - Keep the same <ROW x CELL y> markers in the same order.
        - Return ONLY the table text with the same markers (no extra commentary or explanation).
        - Do not alter the HEADER text.
        
        Begin now:
        """.format('\n'.join(header_paragraphs), table_text_combined, job_description)

    # 5) Send prompt to OpenAI
    completion = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )

    response = completion.choices[0].message.content.strip()

    # 6) Parse the updated text. We expect the LLM to preserve <ROW r CELL c> markers.
    lines = response.splitlines()

    current_row = None
    current_cell = None
    buffer_for_cell = []

    # We'll store the updated text by row/cell in a dict for easy lookup:
    # updated_cells[(row, cell)] = "new text"
    updated_cells = {}

    for line in lines:
        line_stripped = line.strip()
        if line_stripped.startswith("<ROW") and "CELL" in line_stripped:
            # We hit a new marker — meaning we finished the previous cell (if any).
            # 1) Store what we have so far in updated_cells
            if current_row is not None and current_cell is not None:
                updated_cells[(current_row, current_cell)] = "\n".join(buffer_for_cell).strip()

            buffer_for_cell = []  # reset

            # 2) Parse out row/cell indices
            #    Format: <ROW r CELL c>
            #    We'll do a simple parse approach:
            row_part = line_stripped.split("ROW")[1].split("CELL")[0].strip()
            cell_part = line_stripped.split("CELL")[1].split(">")[0].strip()

            current_row = int(row_part)
            current_cell = int(cell_part)

        else:
            # It's actual text content for the current cell
            buffer_for_cell.append(line)

    # Don't forget to store the last cell after loop finishes
    if current_row is not None and current_cell is not None:
        updated_cells[(current_row, current_cell)] = "\n".join(buffer_for_cell).strip()

    # 7) Re-insert the updated text into the table
    for (r, c), new_text in updated_cells.items():
        table_to_update.rows[r].cells[c].text = new_text

    # 8) Generate an output file name and save
    base_name = os.path.splitext(os.path.basename(docx_file))[0]
    output_filename = f"{base_name}_{employer_name}_{job_title}.docx".replace(" ", "_")

    output_path = os.path.join(output_folder, output_filename)
    doc.save(output_path)

    return output_path


# Example Usage
if __name__ == "__main__":
    # Inputs
    docx_file = 'Richard_Raithel_Resume_AI_Engineer.docx'
    job_description_file = "job_desc.txt"
    employer_name = "Raft"
    job_title = "AI Engineer"
    output_folder = os.path.dirname(os.path.abspath(__file__))

    # Read job description from the file
    with open(job_description_file, "r") as f:
        job_description = f.read()

    # Update the resume
    updated_resume_path = update_resume(docx_file, job_description, employer_name, job_title, output_folder)
    print(f"Updated resume saved to: {updated_resume_path}")
